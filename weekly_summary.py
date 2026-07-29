"""Weekly summary Word document.

Rolls the trailing week of daily digests (docs/digest_log.json) and the
active storylines (docs/chronicle.json) into a styled .docx saved under
docs/weekly/, for pasting into the office's monthly email.

Scheduled Friday evenings; a once-per-week output check keeps the backup
cron from duplicating. Manual runs (workflow_dispatch) always regenerate.

Usage:  python weekly_summary.py
"""

import json
import os
import sys
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo

from fedwatch import summarize

WEEK_DAYS = 7  # trailing window; a Friday run covers Sat..Fri (Mon-Fri filled)

EMORY_BLUE = (0x01, 0x21, 0x69)
EMORY_GOLD = (0xF2, 0xA9, 0x00)
EMORY_GRAY = (0x6D, 0x6E, 0x71)


def _load(path, default):
    try:
        with open(path) as f:
            data = json.load(f)
        return data if isinstance(data, type(default)) else default
    except (FileNotFoundError, ValueError):
        return default


def week_window(today=None):
    end = today or datetime.now(ZoneInfo("America/New_York")).date()
    start = end - timedelta(days=WEEK_DAYS - 1)
    return start.strftime("%Y-%m-%d"), end.strftime("%Y-%m-%d")


def week_entries(log, start, end):
    return sorted((d for d in log if start <= (d.get("date") or "") <= end),
                  key=lambda d: d.get("date") or "")


def active_storylines(chronicle, start, end):
    """Storylines with at least one event in the window, most active first."""
    out = []
    for key, s in chronicle.items():
        hits = [e for e in s.get("events", []) if start <= (e.get("date") or "") <= end]
        if hits:
            out.append((len(hits), s))
    return [s for _, s in sorted(out, key=lambda t: -t[0])]


def weekly_narrative(days, storylines, start, end):
    """A 250-350 word weekly overview via Claude; template fallback."""
    if not summarize.claude_available():
        return _template_narrative(days, start, end), "template"
    try:
        import anthropic
        client = anthropic.Anthropic()
        day_block = ""
        for d in days:
            day_block += f"\n{d.get('date')}:\n{(d.get('summary') or '')[:900]}\n"
        story_block = "\n".join(
            f"- {s.get('title')}: {(s.get('summary') or '')[:300]}" for s in storylines)
        prompt = (
            "Write the weekly federal research-policy summary for Emory's Office "
            f"of the SVPR covering {start} to {end}. It will be pasted into a "
            "monthly email to research leadership, so write clean standalone "
            "prose: a 2-3 sentence bottom line for the week, then 'Key "
            "developments' as short bullets grouped by theme (not by day), then "
            "a 1-2 sentence 'What to watch'. 250-350 words, no hype, plain "
            "professional tone. Ground it ONLY in the material below.\n\n"
            f"DAILY SUMMARIES:{day_block}\n\nACTIVE STORYLINES:\n{story_block}"
        )
        resp = client.messages.create(
            model=summarize.MODEL, max_tokens=1200,
            messages=[{"role": "user", "content": prompt}])
        if resp.stop_reason == "refusal":
            return _template_narrative(days, start, end), "template"
        text = next((b.text for b in resp.content if b.type == "text"), "").strip()
        return (text, "claude") if text else (_template_narrative(days, start, end), "template")
    except Exception:  # noqa: BLE001
        return _template_narrative(days, start, end), "template"


def _template_narrative(days, start, end):
    n_items = sum(len(d.get("items", [])) for d in days)
    n_press = sum(len(d.get("press", [])) for d in days)
    return (f"Week of {start} to {end}: {len(days)} digest day(s), {n_items} "
            f"federal actions and {n_press} press-reported developments. See the "
            "item list below; AI narrative was unavailable for this edition.")


def _add_md_text(doc, text):
    """Render the narrative's light markdown into real Word formatting:
    '#' lines become bold headings, '- ' lines become bullets, **bold**
    becomes bold runs. Everything else is a plain paragraph."""
    import re
    from docx.shared import RGBColor
    blue = RGBColor(*EMORY_BLUE)

    def runs(p, line):
        for part in re.split(r"(\*\*[^*]+\*\*)", line):
            if part.startswith("**") and part.endswith("**"):
                p.add_run(part[2:-2]).bold = True
            elif part:
                p.add_run(part.replace("*", ""))

    for raw in (text or "").split("\n"):
        line = raw.strip()
        if not line:
            continue
        if line.startswith("#"):
            p = doc.add_paragraph()
            r = p.add_run(line.lstrip("# ").strip())
            r.bold = True
            r.font.color.rgb = blue
        elif line.startswith(("- ", "* ")):
            runs(doc.add_paragraph(style="List Bullet"), line[2:])
        else:
            runs(doc.add_paragraph(), line)


def _add_hyperlink(paragraph, url, text):
    """python-docx has no hyperlink API; standard OOXML recipe."""
    import docx.opc.constants
    import docx.oxml.ns as ns
    from docx.oxml import OxmlElement
    r_id = paragraph.part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(ns.qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(ns.qn("w:val"), "007dba")
    u = OxmlElement("w:u")
    u.set(ns.qn("w:val"), "single")
    rpr.append(color)
    rpr.append(u)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def build_docx(path, start, end, narrative, days, storylines):
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    blue, gray = RGBColor(*EMORY_BLUE), RGBColor(*EMORY_GRAY)

    title = doc.add_heading("FedWatch Weekly Summary", level=0)
    title.runs[0].font.color.rgb = blue
    sub = doc.add_paragraph(f"Federal research policy · {start} to {end} · "
                            "Office of the Senior Vice President for Research")
    sub.runs[0].font.color.rgb = gray
    sub.runs[0].font.size = Pt(10)

    h = doc.add_heading("Executive overview", level=1)
    h.runs[0].font.color.rgb = blue
    _add_md_text(doc, narrative)

    if storylines:
        h = doc.add_heading("Ongoing storylines", level=1)
        h.runs[0].font.color.rgb = blue
        for s in storylines:
            p = doc.add_paragraph()
            r = p.add_run(s.get("title", ""))
            r.bold = True
            r.font.color.rgb = blue
            if s.get("summary"):
                doc.add_paragraph(s["summary"])

    h = doc.add_heading("This week's items", level=1)
    h.runs[0].font.color.rgb = blue
    for d in days:
        p = doc.add_paragraph()
        r = p.add_run(d.get("date", ""))
        r.bold = True
        r.font.color.rgb = gray
        for it in d.get("items", []):
            p = doc.add_paragraph(style="List Bullet")
            lvl = it.get("level") or ""
            if lvl:
                p.add_run(f"[{lvl}] ").bold = True
            if it.get("url"):
                _add_hyperlink(p, it["url"], it.get("title", ""))
            else:
                p.add_run(it.get("title", ""))
            p.add_run(f"  ({it.get('agency', '')})").font.color.rgb = gray
        for it in d.get("press", []):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("[press] ").bold = True
            if it.get("url"):
                _add_hyperlink(p, it["url"], it.get("title", ""))
            else:
                p.add_run(it.get("title", ""))
            p.add_run(f"  ({it.get('agency', '')})").font.color.rgb = gray

    tail = doc.add_paragraph("Generated by FedWatch. Full daily digests: "
                             "https://chaps72.github.io/ci4hackathon/  ·  "
                             "Historical record: "
                             "https://chaps72.github.io/ci4hackathon/chronicle.html")
    tail.runs[0].font.size = Pt(8)
    tail.runs[0].font.color.rgb = gray
    doc.save(path)


def main() -> int:
    now_et = datetime.now(ZoneInfo("America/New_York"))
    scheduled = os.environ.get("GITHUB_EVENT_NAME", "") == "schedule"
    if scheduled and now_et.weekday() != 4:
        print(f"SKIPPED: weekly summary runs Fridays (today is {now_et:%A}).")
        return 0
    start, end = week_window(now_et.date())
    import pathlib
    out_dir = pathlib.Path("docs/weekly")
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / f"FedWatch-Weekly-{end}.docx"
    if scheduled and out.exists():
        print(f"{out.name} already generated; skipping duplicate cron firing.")
        return 0

    log = _load("docs/digest_log.json", [])
    chronicle = _load("docs/chronicle.json", {})
    days = week_entries(log, start, end)
    if not days:
        print("No digest days in the window; nothing to summarize.")
        return 0
    stories = active_storylines(chronicle, start, end)
    narrative, engine = weekly_narrative(days, stories, start, end)
    build_docx(str(out), start, end, narrative, days, stories)
    print(f"Weekly summary written: {out} ({engine} narrative, "
          f"{len(days)} day(s), {len(stories)} storyline(s)).")

    slack = os.environ.get("SLACK_WEBHOOK_URL", "")
    if slack:
        from fedwatch import notify
        url = f"https://chaps72.github.io/ci4hackathon/weekly/{out.name}"
        notify.send_slack(
            slack, f"This week's Word summary is ready — download: {url}",
            title=f"📄 FedWatch Weekly Summary ({start} to {end})")
        print("Slack: weekly summary link posted.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
